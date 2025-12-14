import logging
from pathlib import Path
from typing import List, Tuple
import chardet
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.core.config import settings
from app.services.interfaces.chunking import IChunkingService
from app.core.enums import DocumentTypeFile
from app.schemas.document import DocumentMetadata, TextChunk, ChunkMetadata

logger = logging.getLogger(__name__)


class ChunkingService(IChunkingService):
    """
    Service for extracting text from documents and splitting into chunks.
    """

    def __init__(
        self,
        chunk_size: int = None,
        chunk_overlap: int = None,
        separators: List[str] = None,
    ):
        """
        Initialize chunking service with LangChain's RecursiveCharacterTextSplitter.
        """
        self._chunk_size = chunk_size or settings.CHUNK_SIZE
        self._chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        self._separators = separators or ["\n\n", "\n", ". ", " ", ""]

        # Initialize LangChain splitter
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=self._chunk_size,
            chunk_overlap=self._chunk_overlap,
            separators=self._separators,
            length_function=len,
            is_separator_regex=False,
        )

        logger.debug(
            "ChunkingService initialized with LangChain splitter",
            extra={
                "chunk_size": self._chunk_size,
                "chunk_overlap": self._chunk_overlap,
                "separators_count": len(self._separators),
            },
        )

    # region TEXT EXTRACTION

    async def extract_text(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """
        Extract text from a file based on its type.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower().lstrip(".")

        if extension not in DocumentTypeFile._value2member_map_:
            raise ValueError(f"Unsupported file extension: {extension}")

        logger.debug(f"Extracting text from {path.name} ({extension})")

        extractor = self._get_extractor(DocumentTypeFile(extension))
        text, metadata = await extractor(file_path)

        # Add common metadata
        metadata.file_name = path.name
        metadata.file_type = extension
        metadata.file_size = path.stat().st_size
        metadata.char_count = len(text)

        logger.debug(
            f"Extracted {len(text)} characters from {path.name}",
            extra={"pages": metadata.pages, "char_count": len(text)},
        )

        return text, metadata

    def _get_extractor(self, extension: DocumentTypeFile) -> callable:
        """Get extraction method based on file extension."""

        match extension:
            case DocumentTypeFile.PDF:
                return self._extract_from_pdf
            case DocumentTypeFile.DOCX:
                return self._extract_from_docx
            case DocumentTypeFile.TXT | DocumentTypeFile.MD:
                return self._extract_from_text
            case _:
                raise ValueError(f"Unsupported file extension: {extension}")

    async def _extract_from_pdf(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from PDF."""

        try:
            reader = PdfReader(file_path)
            pages_text = []
            metadata = DocumentMetadata(pages=len(reader.pages))

            # Extract document info if available
            if reader.metadata:
                if reader.metadata.title:
                    metadata.title = reader.metadata.title
                if reader.metadata.author:
                    metadata.author = reader.metadata.author

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)

            full_text = "\n\n".join(pages_text)
            return full_text, metadata

        except Exception as e:
            logger.error(f"Failed to extract PDF: {e}")
            raise

    async def _extract_from_docx(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from DOCX."""

        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            metadata = DocumentMetadata()

            # Extract document info if available
            if doc.core_properties:
                if doc.core_properties.title:
                    metadata.title = doc.core_properties.title
                if doc.core_properties.author:
                    metadata.author = doc.core_properties.author

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)
            metadata.paragraphs = len(paragraphs)

            return full_text, metadata

        except Exception as e:
            logger.error(f"Failed to extract DOCX: {e}")
            raise

    async def _extract_from_text(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from TXT/MD files."""

        try:
            # Detect encoding
            with open(file_path, "rb") as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected.get("encoding", "utf-8") or "utf-8"

            # Read with detected encoding
            with open(file_path, "r", encoding=encoding, errors="replace") as f:
                text = f.read()

            metadata = DocumentMetadata(
                encoding=encoding,
                lines=text.count("\n") + 1,
            )

            return text, metadata

        except Exception as e:
            logger.error(f"Failed to extract text file: {e}")
            raise

    # endregion TEXT EXTRACTION

    # region TEXT SPLITTING

    async def split_text(
        self,
        text: str,
        metadata: DocumentMetadata = None,
    ) -> List[TextChunk]:
        """
        Split text into chunks using LangChain's RecursiveCharacterTextSplitter.
        """
        if not text or not text.strip():
            return []

        # LangChain splitter to split text into chunks
        chunks = self._splitter.split_text(text)

        # Build chunk objects with metadata
        result = []
        for i, chunk_text in enumerate(chunks):
            chunk_meta = ChunkMetadata(
                chunk_index=i,
                chunk_length=len(chunk_text),
                file_name=metadata.file_name if metadata else None,
                file_type=metadata.file_type if metadata else None,
                pages=metadata.pages if metadata else None,
            )

            result.append(TextChunk(text=chunk_text, metadata=chunk_meta))

        logger.debug(
            f"Split text into {len(result)} chunks",
            extra={
                "total_chars": len(text),
                "avg_chunk_size": sum(c.metadata.chunk_length for c in result) // len(result) if result else 0,
            },
        )

        return result

    # endregion TEXT SPLITTING

    # region CONVENIENCE METHODS

    async def process_file(
        self,
        file_path: str,
    ) -> Tuple[List[TextChunk], DocumentMetadata]:
        """
        Process a file: extract text and split into chunks.
        """

        # Extract text
        text, metadata = await self.extract_text(file_path)

        # Split into chunks
        chunks = await self.split_text(text, metadata)

        return chunks, metadata

    # endregion CONVENIENCE METHODS

    @property
    def chunk_size(self) -> int:
        """Get configured chunk size."""
        return self._chunk_size

    @property
    def chunk_overlap(self) -> int:
        """Get configured chunk overlap."""
        return self._chunk_overlap
